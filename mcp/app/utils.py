"""Formatter for file system items."""
import io
from pypdf import PdfReader
from docx import Document
from PIL import Image as PILImage

from .config import Config

def extract_text_from_file(file_name: str, file_contents: bytes) -> str:
    """
    Extract text from a file. Is used to convert non plain text files to a text.
    Supports: pdf, docx
    """

    if file_name.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(file_contents))
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    elif file_name.endswith(".docx"):
        doc = Document(io.BytesIO(file_contents))
        text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        return '\n'.join(text)
    else:
        return file_contents.decode("utf-8", errors="replace")

def get_image_thumb(image_data: bytes, thumb_width: int, img_format: str) -> bytes:
    """
    Get a thumbnail of the image with the specified width.
    """
    with PILImage.open(io.BytesIO(image_data)) as img:
        img.thumbnail((thumb_width, thumb_width))
        thumb_io = io.BytesIO()
        img.save(thumb_io, format=img_format)
        return thumb_io.getvalue()

def verify_length_is_not_too_large_to_return(contents_length: int, config: Config) -> None:
    """
    Verify that the length of the file contents is not too large.
    """
    if config.max_return_file_size and contents_length > config.max_return_file_size:
        raise ValueError(f"File contents are too large (max: {config.max_return_file_size}, actual: {contents_length})")
